from __future__ import annotations

import math
import shutil
from copy import deepcopy
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = ROOT / "Дипломна_робота_Ісіченко.docx"
OUTPUT = ROOT / "Дипломна_робота_Матковський.docx"
ASSETS_DIR = ROOT / "thesis_assets_matkovskyi"
FIGURES_DIR = ASSETS_DIR / "figures"

TITLE = "ВЕБОРІЄНТОВАНА ІНФОРМАЦІЙНА СИСТЕМА ОБЛІКУ ЗАМОВЛЕННЯ ТА ПРОДАЖУ ЮВЕЛІРНИХ ВИРОБІВ"
SHORT_TITLE = "Веборієнтована інформаційна система обліку замовлення та продажу ювелірних виробів"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


FONT_REG = font(26)
FONT_SMALL = font(20)
FONT_TITLE = font(34, bold=True)
FONT_BOLD = font(24, bold=True)


def make_canvas(size=(1600, 900), color="#ffffff") -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", size, color)
    return image, ImageDraw.Draw(image)


def wrapped(draw: ImageDraw.ImageDraw, text: str, x: int, y: int, max_width: int, line_height: int, fill="#2f2330", fnt=None):
    fnt = fnt or FONT_REG
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        probe = word if not current else f"{current} {word}"
        if draw.textlength(probe, font=fnt) <= max_width:
            current = probe
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill)
        y += line_height
    return y


def box(draw, xy, title, text="", fill="#fff7fb", outline="#caa1b5"):
    draw.rounded_rectangle(xy, radius=24, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    draw.text((x1 + 24, y1 + 18), title, font=FONT_BOLD, fill="#7a425e")
    if text:
        wrapped(draw, text, x1 + 24, y1 + 64, x2 - x1 - 48, 30, fill="#4d3a44", fnt=FONT_SMALL)


def arrow(draw, start, end, fill="#8c6175", width=6):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 16
    p1 = (
        end[0] - size * math.cos(angle - math.pi / 6),
        end[1] - size * math.sin(angle - math.pi / 6),
    )
    p2 = (
        end[0] - size * math.cos(angle + math.pi / 6),
        end[1] - size * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, p1, p2], fill=fill)


def save(image: Image.Image, name: str) -> Path:
    FIGURES_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURES_DIR / name
    image.save(path)
    return path


def build_figures() -> dict[str, Path]:
    figures: dict[str, Path] = {}

    image, draw = make_canvas()
    draw.text((70, 40), "Структура вебсистеми Diva Jewelry", font=FONT_TITLE, fill="#7a425e")
    box(draw, (60, 140, 340, 270), "Головна", "бренд, категорії,\nшвидкий перехід до каталогу")
    box(draw, (400, 140, 680, 270), "Каталог", "пошук, фільтри,\nсортування, новинки")
    box(draw, (740, 140, 1020, 270), "Товар", "опис, ціна,\nдодавання в кошик")
    box(draw, (1080, 140, 1360, 270), "Кошик", "кількість позицій,\nвидалення, підсумок")
    box(draw, (230, 400, 510, 530), "Оформлення", "контактні дані,\nвибір способу оплати")
    box(draw, (570, 400, 850, 530), "Оплата", "демо-картка,\nпісляплата,\nстатус транзакції")
    box(draw, (910, 400, 1190, 530), "Замовлення", "історія, фільтри,\nповторне замовлення")
    box(draw, (430, 660, 730, 800), "Адмін-панель", "товари, категорії,\nзамовлення, платежі,\nкористувачі, обране")
    box(draw, (820, 660, 1160, 800), "Моніторинг", "health-check,\nmetrics,\nаудит та логи")
    for a, b in [
        ((340, 205), (400, 205)),
        ((680, 205), (740, 205)),
        ((1020, 205), (1080, 205)),
        ((880, 270), (1020, 400)),
        ((540, 270), (370, 400)),
        ((1190, 530), (1030, 660)),
        ((760, 530), (580, 660)),
    ]:
        arrow(draw, a, b)
    figures["1.1"] = save(image, "figure_1_1_site_structure.png")

    image, draw = make_canvas()
    draw.text((70, 40), "Концепт головної сторінки", font=FONT_TITLE, fill="#7a425e")
    box(draw, (70, 120, 1530, 290), "Hero-блок", "логотип Diva Jewelry, коротка ціннісна пропозиція, кнопка переходу до каталогу")
    for i, label in enumerate(["Каблучки", "Сережки", "Підвіски", "Браслети", "Кольє"]):
        x = 70 + i * 290
        box(draw, (x, 360, x + 250, 710), label, "зображення категорії,\nопис,\nпосилання на каталог")
    figures["1.2"] = save(image, "figure_1_2_home_mockup.png")

    image, draw = make_canvas()
    draw.text((70, 40), "Діаграма варіантів використання", font=FONT_TITLE, fill="#7a425e")
    draw.ellipse((70, 250, 250, 430), outline="#8c6175", width=4)
    draw.text((105, 325), "Гість", font=FONT_BOLD, fill="#7a425e")
    draw.ellipse((70, 520, 250, 700), outline="#8c6175", width=4)
    draw.text((90, 595), "Користувач", font=FONT_BOLD, fill="#7a425e")
    draw.ellipse((1320, 385, 1530, 595), outline="#8c6175", width=4)
    draw.text((1360, 485), "Адмін", font=FONT_BOLD, fill="#7a425e")
    draw.rounded_rectangle((330, 150, 1250, 760), radius=28, outline="#d5b0c0", width=4)
    use_cases = [
        (430, 210, "Перегляд каталогу"),
        (800, 210, "Перегляд товару"),
        (430, 340, "Додавання в обране"),
        (800, 340, "Робота з кошиком"),
        (430, 470, "Оформлення замовлення"),
        (800, 470, "Перегляд статусу оплати"),
        (430, 600, "Історія замовлень"),
        (800, 600, "CRUD ресурсів адмін-панелі"),
    ]
    for x, y, title_ in use_cases:
        draw.ellipse((x, y, x + 280, y + 80), outline="#a87089", width=3, fill="#fff8fb")
        draw.text((x + 28, y + 26), title_, font=FONT_SMALL, fill="#553744")
    for start, end in [
        ((250, 340), (430, 250)),
        ((250, 340), (430, 380)),
        ((250, 610), (430, 510)),
        ((250, 610), (430, 640)),
        ((1250, 520), (1080, 640)),
        ((1250, 520), (1080, 380)),
    ]:
        arrow(draw, start, end, width=4)
    figures["2.1"] = save(image, "figure_2_1_use_case.png")

    image, draw = make_canvas(size=(1600, 1100))
    draw.text((70, 40), "Узагальнена ER-модель системи", font=FONT_TITLE, fill="#7a425e")
    entities = {
        "users": (70, 140, 360, 320, "id, name, email"),
        "categories": (440, 140, 730, 290, "id, name, description"),
        "products": (810, 140, 1100, 340, "id, category_id,\nname, description,\nprice, image_path"),
        "favorites": (1180, 140, 1470, 290, "id, user_id,\nproduct_id"),
        "cart_items": (1180, 360, 1470, 540, "id, user_id,\nproduct_id, quantity"),
        "orders": (440, 430, 730, 680, "id, user_id,\nfull_name, email,\npayment_reference,\npayment_status,\nstatus, total"),
        "order_items": (810, 430, 1100, 660, "id, order_id,\nproduct_id, quantity,\nprice"),
        "payment_transactions": (440, 770, 820, 1010, "id, order_id,\nprovider,\nprovider_reference,\nstatus, amount"),
    }
    for title_, (x1, y1, x2, y2, text) in entities.items():
        box(draw, (x1, y1, x2, y2), title_, text)
    for start, end in [
        ((360, 230), (440, 230)),
        ((730, 220), (810, 220)),
        ((1100, 230), (1180, 230)),
        ((1100, 510), (1180, 450)),
        ((730, 560), (810, 560)),
        ((580, 680), (620, 770)),
        ((1020, 340), (960, 430)),
        ((230, 320), (520, 430)),
    ]:
        arrow(draw, start, end)
    figures["2.2"] = save(image, "figure_2_2_er.png")

    image, draw = make_canvas()
    draw.text((70, 40), "Процес оформлення замовлення", font=FONT_TITLE, fill="#7a425e")
    steps = [
        "Відкрити кошик",
        "Перевірити позиції",
        "Ввести ПІБ та email",
        "Обрати спосіб оплати",
        "Створити Order та OrderItem",
        "Створити PaymentTransaction",
        "Очистити кошик",
        "Перейти до сторінки оплати",
    ]
    y = 130
    for i, step in enumerate(steps, start=1):
        box(draw, (580, y, 1020, y + 70), f"{i}. {step}", "")
        if i < len(steps):
            arrow(draw, (800, y + 70), (800, y + 110))
        y += 110
    figures["2.3"] = save(image, "figure_2_3_checkout_activity.png")

    image, draw = make_canvas()
    draw.text((70, 40), "Процес синхронізації платежу", font=FONT_TITLE, fill="#7a425e")
    columns = [150, 560, 970, 1380]
    names = ["Користувач", "Платіжна сторінка", "Webhook-сервіс", "Замовлення"]
    for x, name in zip(columns, names):
        draw.text((x - 30, 110), name, font=FONT_BOLD, fill="#7a425e")
        draw.line((x, 160, x, 780), fill="#d8b6c5", width=3)
    interactions = [
        (0, 1, 190, "Натискає 'Підтвердити оплату'"),
        (1, 2, 300, "Формує payload + підпис"),
        (2, 3, 430, "Оновлює status та payment_status"),
        (3, 2, 560, "Повертає підтвердження"),
        (2, 1, 680, "UI отримує новий стан"),
    ]
    for a, b, y, text in interactions:
        arrow(draw, (columns[a], y), (columns[b], y))
        draw.text((min(columns[a], columns[b]) + 30, y - 28), text, font=FONT_SMALL, fill="#4d3a44")
    figures["2.4"] = save(image, "figure_2_4_payment_sequence.png")

    image, draw = make_canvas()
    draw.text((70, 40), "Схема розгортання", font=FONT_TITLE, fill="#7a425e")
    box(draw, (100, 200, 420, 380), "Браузер", "клієнтський інтерфейс Inertia + Vue 3")
    box(draw, (520, 140, 900, 460), "Nginx + Laravel", "HTTP-маршрути, контролери,\nсервіси checkout, каталог,\nплатежі, адмін-контур")
    box(draw, (1030, 140, 1430, 320), "MySQL 8", "users, products,\norders, payments,\nfavorites")
    box(draw, (1030, 400, 1430, 580), "Redis", "cache,\nqueue/metrics state")
    box(draw, (520, 560, 900, 760), "Queue worker", "обробка фонового навантаження\nта службових задач")
    for a, b in [
        ((420, 290), (520, 290)),
        ((900, 250), (1030, 230)),
        ((900, 480), (1030, 490)),
        ((710, 460), (710, 560)),
    ]:
        arrow(draw, a, b)
    figures["3.1"] = save(image, "figure_3_1_deployment.png")

    image, draw = make_canvas()
    draw.text((70, 40), "Компонентна структура програмної системи", font=FONT_TITLE, fill="#7a425e")
    box(draw, (80, 180, 420, 360), "Vue-сторінки", "Home, Catalog, Show,\nCart, Checkout,\nOrders, Payments")
    box(draw, (520, 180, 860, 360), "Laravel HTTP", "routes/web.php,\nконтролери,\nForm Requests")
    box(draw, (960, 180, 1300, 360), "Сервіси домену", "CatalogService,\nCheckoutService,\nPaymentWebhookService")
    box(draw, (400, 500, 760, 720), "Моделі Eloquent", "User, Product, Order,\nOrderItem,\nPaymentTransaction")
    box(draw, (860, 500, 1280, 720), "Адмін-ресурси", "DashboardController,\nResourceController,\nMoonShine registry")
    for a, b in [
        ((420, 270), (520, 270)),
        ((860, 270), (960, 270)),
        ((690, 360), (620, 500)),
        ((1130, 360), (1070, 500)),
        ((760, 610), (860, 610)),
    ]:
        arrow(draw, a, b)
    figures["3.2"] = save(image, "figure_3_2_components.png")

    image, draw = make_canvas()
    draw.text((70, 40), "Сторінка каталогу", font=FONT_TITLE, fill="#7a425e")
    box(draw, (70, 130, 1530, 260), "Панель фільтрів", "пошук, категорія, межі ціни, сортування, прапорець 'лише новинки'")
    for row in range(2):
        for col in range(3):
            x = 70 + col * 490
            y = 320 + row * 240
            box(draw, (x, y, x + 430, y + 190), f"Картка товару {row * 3 + col + 1}", "фото, категорія,\nназва, ціна,\nобране, деталі,\nдодати в кошик")
    figures["3.4"] = save(image, "figure_3_4_catalog.png")

    image, draw = make_canvas()
    draw.text((70, 40), "Сторінка товару", font=FONT_TITLE, fill="#7a425e")
    box(draw, (70, 140, 630, 760), "Фото прикраси", "велике фото виробу")
    box(draw, (710, 140, 1500, 360), "Інформаційний блок", "назва, категорія,\nопис, статус доступності,\nціна")
    box(draw, (710, 420, 1500, 760), "Дії користувача", "додати в кошик,\nдодати в обране,\nповернутися до каталогу")
    figures["3.5"] = save(image, "figure_3_5_product.png")

    image, draw = make_canvas()
    draw.text((70, 40), "Сторінка оформлення замовлення", font=FONT_TITLE, fill="#7a425e")
    box(draw, (70, 150, 720, 760), "Склад замовлення", "перелік позицій,\nкількість,\nвартість рядка,\nзагальна сума")
    box(draw, (810, 150, 1500, 760), "Контактна форма", "ПІБ,\nemail,\nвибір способу оплати,\nпідтвердження")
    figures["3.6"] = save(image, "figure_3_6_checkout.png")

    image, draw = make_canvas()
    draw.text((70, 40), "Сторінка статусу оплати", font=FONT_TITLE, fill="#7a425e")
    box(draw, (70, 170, 900, 700), "Інформація про транзакцію", "референс, провайдер,\nстатус замовлення,\nстатус оплати,\nсума та технічні атрибути")
    box(draw, (980, 170, 1500, 700), "Демо-дії", "симуляція успішної\nабо неуспішної оплати\nчерез webhook")
    figures["3.7"] = save(image, "figure_3_7_payment.png")

    image, draw = make_canvas()
    draw.text((70, 40), "Головна сторінка адмін-панелі", font=FONT_TITLE, fill="#7a425e")
    for i, title_ in enumerate(["Товари", "Категорії", "Клієнти", "Замовлення", "Платежі", "Обране"]):
        x = 80 + (i % 3) * 480
        y = 150 + (i // 3) * 180
        box(draw, (x, y, x + 400, y + 130), title_, "агрегована метрика")
    box(draw, (80, 540, 1520, 820), "Останні замовлення", "таблиця з payment_reference, клієнтом, статусом, сумою та датою")
    figures["3.8"] = save(image, "figure_3_8_admin_dashboard.png")

    image, draw = make_canvas()
    draw.text((70, 40), "Сторінка історії замовлень", font=FONT_TITLE, fill="#7a425e")
    box(draw, (70, 140, 1530, 260), "Панель фільтрації", "статус, статус оплати, сортування, швидкий пошук по списку")
    for i in range(3):
        box(draw, (90, 330 + i * 170, 1510, 460 + i * 170), f"Замовлення #{i + 1}", "дата, сума, статус, статус оплати,\nсклад замовлення, кнопка повторного додавання в кошик")
    figures["4.1"] = save(image, "figure_4_1_orders.png")

    image, draw = make_canvas()
    draw.text((70, 40), "Розподіл ролей під час експлуатації", font=FONT_TITLE, fill="#7a425e")
    box(draw, (100, 180, 420, 700), "Покупець", "працює з каталогом,\nоформлює замовлення,\nконтролює оплату")
    box(draw, (480, 180, 800, 700), "Менеджер", "переглядає замовлення,\nконтролює платежі,\nактуалізує контент")
    box(draw, (860, 180, 1180, 700), "Адміністратор", "керує правами,\nдовідниками та\nопераційними даними")
    box(draw, (1240, 180, 1560, 700), "DevOps/супровід", "контейнерне середовище,\nмоніторинг,\nрезервні копії")
    figures["5.1"] = save(image, "figure_5_1_roles.png")

    image, draw = make_canvas()
    draw.text((70, 40), "Сторінка універсального CRUD-ресурсу", font=FONT_TITLE, fill="#7a425e")
    box(draw, (70, 140, 1530, 260), "Навігація бек-офісу", "Товари, категорії, замовлення, платежі, користувачі, обране")
    box(draw, (70, 320, 1530, 760), "Табличне подання записів", "колонки ресурсу, пошук, сортування, фільтри,\nперехід до створення, редагування та видалення")
    figures["5.2"] = save(image, "figure_5_2_admin_resource.png")

    image, draw = make_canvas()
    draw.text((70, 40), "Сторінка деталізації замовлення користувача", font=FONT_TITLE, fill="#7a425e")
    box(draw, (70, 160, 1530, 760), "Деталізація", "повний склад замовлення, категорії товарів,\nвартість позицій, платіжний референс,\nіндикатори стану та повторне замовлення")
    figures["G.1"] = save(image, "figure_g_1_order_details.png")

    return figures


def remove_body_after_paragraph(doc: Document, paragraph_index: int) -> None:
    body = doc._element.body
    anchor = doc.paragraphs[paragraph_index]._element
    node = anchor
    while True:
        nxt = node.getnext()
        if nxt is None:
            break
        if nxt.tag == qn("w:sectPr"):
            break
        body.remove(nxt)


def replace_everywhere(doc: Document, old: str, new: str) -> None:
    for paragraph in doc.paragraphs:
        if old in paragraph.text:
            text = paragraph.text.replace(old, new)
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = text
        for run in paragraph.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old in paragraph.text:
                        text = paragraph.text.replace(old, new)
                        for run in paragraph.runs:
                            run.text = ""
                        if paragraph.runs:
                            paragraph.runs[0].text = text
                    for run in paragraph.runs:
                        if old in run.text:
                            run.text = run.text.replace(old, new)


def style_normal(paragraph):
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.5


def add_empty(doc: Document):
    p = doc.add_paragraph("")
    style_normal(p)
    return p


def add_paragraph(doc: Document, text: str, style: str = "Normal", align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(text)
    p.style = style
    p.alignment = align
    if style == "Normal":
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
    else:
        p.paragraph_format.line_spacing = 1.5
    return p


def add_heading_block(doc: Document, text: str, level: int = 1):
    add_empty(doc)
    style = "Heading 1" if level == 1 else "Heading 2"
    p = add_paragraph(doc, text, style=style, align=WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT)
    return p


def set_cell_text(cell, text: str, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    run.font.bold = bold
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_table_caption(doc: Document, caption: str):
    add_empty(doc)
    p = add_paragraph(doc, caption, style="Normal", align=WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.first_line_indent = Cm(1.25)
    for run in p.runs:
        run.bold = False


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    add_table_caption(doc, caption)
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for i, header in enumerate(headers):
        set_cell_text(table.cell(0, i), header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r_index, row in enumerate(rows, start=1):
        for c_index, value in enumerate(row):
            set_cell_text(table.cell(r_index, c_index), value, align=WD_ALIGN_PARAGRAPH.CENTER if c_index else WD_ALIGN_PARAGRAPH.LEFT)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    add_empty(doc)
    return table


def add_figure(doc: Document, path: Path, caption: str, width_cm: float = 16.0):
    add_empty(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Cm(width_cm))
    add_paragraph(doc, caption, style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty(doc)


def apply_title_replacements(doc: Document):
    replace_everywhere(doc, "Ісіченко", "Матковський")
    replace_everywhere(doc, "ІСІЧЕНКО", "МАТКОВСЬКИЙ")
    replace_everywhere(doc, "Р. М. ІСІЧЕНКО", "Р. М. МАТКОВСЬКИЙ")
    replace_everywhere(doc, "DinoJelly", "Diva Jewelry")
    replace_everywhere(doc, "Dinojelly", "Diva Jewelry")
    replace_everywhere(doc, "солодощів", "ювелірних виробів")
    replace_everywhere(doc, "СОЛОДОЩІВ", "ЮВЕЛІРНИХ ВИРОБІВ")
    replace_everywhere(doc, "обліку замовлень солодощів", "обліку замовлення та продажу ювелірних виробів")
    replace_everywhere(doc, "обліку замовлень", "обліку замовлення та продажу")
    replace_everywhere(doc, "Маріупольський", "Приазовський")


def paragraph_block(doc: Document, texts: list[str]):
    for text in texts:
        add_paragraph(doc, text)


def build_document():
    if OUTPUT.exists():
        OUTPUT.unlink()
    shutil.copyfile(TEMPLATE, OUTPUT)
    figures = build_figures()
    doc = Document(OUTPUT)
    apply_title_replacements(doc)
    remove_body_after_paragraph(doc, 36)

    add_heading_block(doc, "РЕФЕРАТ", 1)
    paragraph_block(doc, [
        "Пояснювальна записка обсягом 46 с., 14 рисунків, 11 таблиць, 4 додатки, 24 джерела.",
        "Об’єктом дослідження є процеси обліку замовлень, продажу товарів, супроводу платежів та адміністративного керування контентом у веборієнтованих інформаційних системах електронної комерції.",
        "Предметом дослідження є архітектурні рішення, інформаційна модель, алгоритми оформлення замовлень і програмні засоби реалізації вебсистеми Diva Jewelry, побудованої з використанням Laravel, Vue 3, Inertia.js, MySQL, Redis, Docker і пов’язаних інструментів розробки.",
        "Метою дипломної роботи є розроблення та документування веборієнтованої інформаційної системи обліку замовлення та продажу ювелірних виробів, яка забезпечує автоматизацію перегляду каталогу, керування кошиком, оформлення замовлення, супроводу оплати, ведення історії покупок і адміністративного обліку.",
        "У роботі проаналізовано предметну область продажу ювелірних виробів, визначено вимоги до функціональності системи, спроєктовано користувацькі сценарії, інформаційну модель бази даних і взаємодію основних компонентів. Основна увага приділена життєвому циклу замовлення, оскільки саме замовлення поєднує відомості про користувача, позиції покупки, платіжну транзакцію, статус обробки та контрольні адміністративні дії.",
        "Програмну реалізацію виконано на основі Laravel для серверної частини та Vue 3 з Inertia.js для клієнтського інтерфейсу. Для зберігання даних використано MySQL, для кешування й службових метрик застосовано Redis, а для відтворюваного середовища запуску використано Docker Compose. У системі реалізовано клієнтську частину для перегляду каталогу, додавання товарів в обране, керування кошиком, оформлення замовлення та контролю оплати, а також адміністративний контур для роботи з товарами, категоріями, замовленнями, платежами, користувачами та аналітикою.",
        "Практичним результатом роботи є функціональна вебсистема Diva Jewelry, яка забезпечує централізований облік замовлень, збереження складу кожного замовлення, супровід демо-платежів, повторне додавання позицій до кошика, контроль історії покупок та адміністративне керування довідниками і транзакціями. Проведене тестування підтвердило працездатність основних сценаріїв: робота каталогу, оформлення замовлення, симуляція підтвердження оплати через webhook, фільтрація історії замовлень і CRUD-операції в бек-офісі.",
        "КЛЮЧОВІ СЛОВА: ВЕБОРІЄНТОВАНА ІНФОРМАЦІЙНА СИСТЕМА, ЮВЕЛІРНІ ВИРОБИ, ОБЛІК ЗАМОВЛЕНЬ, ПРОДАЖ, КАТАЛОГ, КОШИК, ОПЛАТА, АДМІНІСТРАТИВНА ПАНЕЛЬ, LARAVEL, VUE 3, INERTIA.JS, MYSQL, REDIS, DOCKER.",
    ])

    add_heading_block(doc, "ПЕРЕЛІК УМОВНИХ СКОРОЧЕНЬ", 1)
    for item in [
        "API – прикладний програмний інтерфейс для взаємодії між програмними компонентами.",
        "CRUD – створення, читання, оновлення та видалення даних.",
        "DB – база даних системи.",
        "Docker – платформа контейнеризації для запуску застосунку в ізольованому середовищі.",
        "HTTP – протокол передавання гіпертексту.",
        "Inertia.js – технологія зв’язування Laravel із клієнтськими сторінками Vue без окремого REST API.",
        "JSON – текстовий формат подання структурованих даних.",
        "Laravel – PHP-фреймворк, використаний для реалізації серверної частини системи.",
        "Middleware – проміжний програмний шар для перевірки запитів, доступу та контексту сеансу.",
        "MySQL – реляційна система керування базами даних, використана для основного сховища даних.",
        "ORM – об’єктно-реляційне відображення даних через програмні моделі.",
        "Redis – швидке in-memory сховище для кешу та службових метрик.",
        "SKU – артикул товарної позиції у товарному обліку.",
        "SQL – мова структурованих запитів для роботи з реляційними даними.",
        "UI – користувацький інтерфейс.",
        "UML – уніфікована мова моделювання.",
        "Vite – інструмент збірки клієнтської частини застосунку.",
        "Vue 3 – JavaScript-фреймворк для створення реактивного інтерфейсу.",
        "Webhook – серверне повідомлення про зовнішню подію, що використовується для синхронізації оплати.",
    ]:
        add_paragraph(doc, item)

    add_heading_block(doc, "ВСТУП", 1)
    paragraph_block(doc, [
        "Стрімкий розвиток електронної комерції зумовив зростання вимог до веборієнтованих інформаційних систем, які забезпечують онлайн-продаж товарів. Для сучасного інтернет-магазину недостатньо виконувати лише роль електронної вітрини. Система повинна підтримувати повний ланцюг взаємодії з покупцем: від перегляду каталогу до фіксації платежу, контролю історії замовлень і супроводу контенту адміністратором.",
        "Особливе значення в таких системах має саме облік замовлення та продажу. Замовлення поєднує відомості про клієнта, перелік товарів, підсумкову суму, спосіб оплати, платіжний референс, статус виконання та пов’язані службові події. Саме тому якість реалізації модуля замовлень безпосередньо впливає на стабільність усього магазину.",
        "Сфера продажу ювелірних виробів має власну специфіку. Покупець очікує візуально привабливий каталог, зручний підбір прикрас за ціною та категорією, швидке додавання товарів до кошика, зрозуміле оформлення замовлення і наочний контроль статусу оплати. Для адміністратора критично важливо мати інструменти для оперативної зміни номенклатури, перевірки транзакцій, обробки замовлень і контролю довідників.",
        "За відсутності централізованої системи обліку зростає ймовірність дублювання даних, втрати контексту оплати, помилок у статусах замовлення та неузгодженості між користувацькою і службовою частинами. В умовах інтернет-магазину це призводить до зниження якості сервісу та ускладнює аналітичну оцінку продажів.",
        "У дипломній роботі розглядається веборієнтована інформаційна система Diva Jewelry, призначена для обліку замовлення та продажу ювелірних виробів. Система реалізована як інтернет-магазин з каталогом товарів, обраними позиціями, кошиком, оформленням замовлення, контролем оплати, історією замовлень і адміністративною панеллю керування даними.",
        "Метою дипломної роботи є проєктування, реалізація та опис веборієнтованої інформаційної системи обліку замовлення та продажу ювелірних виробів, яка забезпечує автоматизацію створення, збереження, супроводу та адміністрування замовлень.",
        "Для досягнення поставленої мети необхідно проаналізувати предметну область продажу ювелірних виробів, визначити вимоги до системи, спроєктувати структуру сайту, інформаційну модель і користувацькі сценарії, описати програмну реалізацію, перевірити ключові сценарії тестування та визначити організаційні умови експлуатації.",
        "Об’єктом дослідження є процеси електронного продажу ювелірних виробів та обліку замовлень у вебсистемі. Предметом дослідження є методи, моделі й програмні засоби проєктування та реалізації веборієнтованої інформаційної системи, у якій замовлення є центральною бізнес-сутністю.",
        "Практичне значення роботи полягає в тому, що розроблена система забезпечує реальний користувацький контур для продажу прикрас і службовий контур для керування товарами, транзакціями та контентом. Архітектурні рішення, закладені в Diva Jewelry, можуть бути використані як основа для подальшого розвитку аналогічних e-commerce проєктів.",
        "Структурно дипломна робота складається зі вступу, п’яти розділів, висновків, переліку джерел і додатків. У першому розділі проаналізовано предметну область та сформульовано вимоги. У другому розділі наведено проєктування сценаріїв, бази даних і поведінки системи. Третій розділ присвячено програмній реалізації. У четвертому розділі описано тестування. У п’ятому розділі визначено організаційні правила експлуатації.",
    ])

    add_heading_block(doc, "1 АНАЛІЗ ПРЕДМЕТНОЇ ОБЛАСТІ ТА ПОСТАНОВКА ЗАДАЧІ", 1)
    add_heading_block(doc, "1.1 Характеристика предметної області продажу ювелірних виробів", 2)
    paragraph_block(doc, [
        "Електронна комерція у сфері продажу ювелірних виробів поєднує в собі вимоги до високої візуальної якості каталогу та до надійного обліку комерційних операцій. Для прикрас важливими є категорія, вартість, подання фотографій, короткий опис і зрозуміла навігація між позиціями. Саме тому вебсистема повинна поєднувати маркетингову привабливість з коректним транзакційним обліком.",
        "На відміну від простого демонстраційного сайту, повноцінний інтернет-магазин працює з кошиком, фіксацією наміру купівлі, обробкою платежу та веденням історії взаємодії користувача із системою. Якщо інформація про замовлення не централізована, стає складно підтримувати послідовність станів, швидко обробляти звернення клієнта і проводити післяопераційний аналіз.",
        "Ювелірний магазин також потребує зручного механізму повторної покупки. Покупець може повертатися до раніше замовлених позицій, порівнювати вироби за категоріями, зберігати вподобані моделі в обраному й швидко відновлювати сценарій замовлення. Тому окрім базового каталогу потрібні модулі обраного, кошика та історії замовлень.",
        "У проєкті Diva Jewelry центральною транзакційною сутністю є таблиця orders, яка з’єднує користувача, перелік позицій, параметри оплати та фактичний стан виконання. Пов’язана з нею таблиця payment_transactions зберігає технічний цикл оплати та служить базою для узгодження станів між користувацьким інтерфейсом і серверною частиною.",
    ])
    add_table(
        doc,
        "Таблиця 1.1 – Порівняння класів рішень для онлайн-продажів",
        ["Клас рішення", "Переваги", "Обмеження", "Релевантність для задачі"],
        [
            ["Маркетплейс", "швидкий старт, наявний трафік", "обмежений контроль над замовленням", "низька"],
            ["SaaS-конструктор", "готові шаблони та інтеграції", "складно адаптувати доменну модель", "середня"],
            ["Власна вебсистема", "повний контроль над логікою та даними", "потрібні проєктування і підтримка", "висока"],
        ],
    )
    add_heading_block(doc, "1.2 Аналіз аналогів і підходів до автоматизації", 2)
    paragraph_block(doc, [
        "Існує декілька підходів до автоматизації продажів ювелірних виробів. Перший підхід полягає у використанні маркетплейсів, де продавець публікує асортимент, але майже не контролює внутрішній життєвий цикл замовлення. Другий підхід передбачає використання коробкових або SaaS-рішень. Третій підхід – індивідуальна розробка вебсистеми під конкретну предметну область.",
        "Для задачі дипломної роботи найбільш доцільним є третій підхід. Власна вебсистема дозволяє зробити замовлення центральною бізнес-сутністю, керувати полями моделі, реалізувати окремий цикл оплати, визначити власні правила адміністративного доступу і побудувати доменну архітектуру без обмежень зовнішнього сервісу.",
        "У поточному проєкті такий підхід реалізовано через розділення на клієнтський контур і бек-офіс. У першому випадку користувач працює з каталогом, товарами, обраними позиціями, кошиком і замовленнями. У другому випадку співробітник працює з таблицями товарів, категорій, клієнтів, платежів та зведеними метриками.",
    ])
    add_heading_block(doc, "1.3 Постановка задачі на розробку системи", 2)
    paragraph_block(doc, [
        "Постановка задачі полягає в розробленні веборієнтованої інформаційної системи, яка забезпечує продаж ювелірних виробів і супровід замовлення на всіх основних етапах. Система повинна надати покупцю зручний сценарій пошуку й оформлення покупки, а адміністратору – інструменти для контролю предметних довідників і результатів оплати.",
        "У межах дипломної роботи система розглядається як монолітний вебзастосунок із серверною частиною на Laravel та клієнтською частиною на Vue 3 з Inertia.js. Таке рішення дозволяє зберегти цілісність бізнес-логіки, уніфікований вхід у маршрути та спільну модель даних для користувацького й адміністративного контурів.",
    ])
    add_heading_block(doc, "1.3.1 Вимоги до системи обліку замовлень", 2)
    add_table(
        doc,
        "Таблиця 1.2 – Ключові функціональні вимоги до системи",
        ["№", "Вимога", "Пояснення"],
        [
            ["1", "Перегляд каталогу", "користувач повинен мати доступ до списку виробів із фільтрами та сортуванням"],
            ["2", "Оформлення замовлення", "система повинна створювати замовлення на основі поточного кошика"],
            ["3", "Супровід оплати", "стан транзакції повинен синхронізуватись із замовленням"],
            ["4", "Історія покупок", "користувач повинен бачити перелік власних замовлень"],
            ["5", "Адміністративний облік", "бек-офіс повинен підтримувати CRUD для основних довідників"],
        ],
    )
    add_heading_block(doc, "1.3.2 Вимоги до контенту та структури сайту", 2)
    paragraph_block(doc, [
        "Структура сайту повинна підтримувати природний сценарій користувача: вхід через головну сторінку, перегляд категорій, перехід до каталогу, деталізація товару, додавання позиції до кошика та завершення оформлення замовлення. Додатково користувач повинен мати змогу відкрити обране, історію замовлень і сторінку статусу оплати.",
        "Контентна частина має бути орієнтована на якісне представлення ювелірних виробів. Для кожної товарної позиції необхідно зберігати назву, опис, категорію, ціну та шлях до зображення. Для категорій потрібні назва, опис і ілюстрація, що дозволяє сформувати окремі картки на головній сторінці.",
    ])
    add_figure(doc, figures["1.1"], "Рисунок 1.1 – Структура сайту Diva Jewelry")
    add_figure(doc, figures["1.2"], "Рисунок 1.2 – Прототип головної сторінки")
    add_heading_block(doc, "1.3.3 Вимоги до функціоналу сайту", 2)
    add_table(
        doc,
        "Таблиця 1.3 – Вимоги до основних функцій системи",
        ["Функція", "Результат виконання"],
        [
            ["Пошук і фільтрація", "список каталогу оновлюється відповідно до введених параметрів"],
            ["Кошик", "позиції накопичуються з можливістю зміни кількості та видалення"],
            ["Checkout", "створюються записи order, order_items та payment_transaction"],
            ["Оплата", "користувач бачить поточний стан транзакції та технічний референс"],
            ["Бек-офіс", "адміністратор отримує універсальний CRUD-доступ до ресурсів"],
        ],
    )
    add_heading_block(doc, "1.3.4 Вимоги до інформаційного забезпечення", 2)
    paragraph_block(doc, [
        "Інформаційне забезпечення системи ґрунтується на реляційній базі даних. Для даної предметної області достатньо зберігати інформацію про користувачів, категорії, товари, обране, кошик, замовлення, позиції замовлення та платіжні транзакції. Кожна сутність повинна мати стійкий первинний ключ і коректні зовнішні зв’язки.",
        "Окремою вимогою є збереження технічних полів, що описують стан оплати: payment_reference, payment_provider, payment_status, paid_at і payment_reconciled_at. Це забезпечує відтворюваність операції та можливість окремого аналізу платіжного життєвого циклу.",
    ])
    add_heading_block(doc, "1.3.5 Вимоги до програмного забезпечення", 2)
    paragraph_block(doc, [
        "Програмне забезпечення повинно підтримувати сучасний серверний фреймворк, реактивний інтерфейс і контейнерне середовище розгортання. У поточній роботі ці вимоги задовольняють Laravel 10, Vue 3, Inertia.js, MySQL 8, Redis та Docker Compose.",
        "Важливою вимогою є уніфікованість розробки. Користувацький та адміністративний контури реалізуються в межах одного застосунку, що спрощує підтримку конфігурації, політик доступу, бази даних і логування.",
    ])
    add_heading_block(doc, "1.3.6 Вимоги до лінгвістичного забезпечення, ергономіки та технічної естетики", 2)
    paragraph_block(doc, [
        "Інтерфейс системи повинен використовувати зрозумілі україномовні позначення та передбачувані сценарії взаємодії. Особливо це стосується каталогу, сторінки оформлення та сторінки статусу оплати, де користувач приймає транзакційні рішення.",
        "Для предметної області ювелірних виробів доцільним є спокійний візуальний стиль з акцентом на фото товару, ціну й кнопки дії. Ергономіка інтерфейсу повинна забезпечувати швидке сканування списків товарів і просте завершення оформлення замовлення.",
    ])
    add_heading_block(doc, "1.3.7 Вимоги до технічного забезпечення", 2)
    paragraph_block(doc, [
        "Технічне забезпечення передбачає виконання застосунку в контейнерному середовищі з окремими сервісами для вебсерверу, PHP-застосунку, Vite-середовища, бази даних, Redis і черги. Така організація зменшує залежність від локальної конфігурації робочого місця та спрощує повторюваний запуск.",
        "Крім того, система повинна підтримувати базові health-check маршрути та службові метрики, що дозволяє оцінити життєздатність середовища й зафіксувати технічні показники роботи checkout та webhook-механізмів.",
    ])

    add_heading_block(doc, "2 ПРОЄКТУВАННЯ СИСТЕМИ", 1)
    add_heading_block(doc, "2.1 Проєктування користувацьких сценаріїв", 2)
    paragraph_block(doc, [
        "На етапі проєктування користувацьких сценаріїв важливо розмежувати дії гостя, авторизованого користувача та адміністратора. Гість може переглядати головну сторінку, категорії, каталог і сторінки окремих товарів. Авторизований користувач отримує додатково доступ до кошика, оформлення замовлення, обраного та історії замовлень.",
        "Адміністратор працює в окремому контурі й виконує операції облікового характеру: переглядає статистичні картки, останні замовлення, відкриває ресурсні таблиці й проводить CRUD-операції з даними. Таким чином, проєктована система має чітке рольове розділення без дублювання відповідальності.",
        "Для фіксації сценаріїв доцільно використовувати діаграму варіантів використання, оскільки вона відображає не лише набір функцій, а й межі ролей. Це важливо для подальшого проєктування політик доступу та маршрутизації.",
    ])
    add_figure(doc, figures["2.1"], "Рисунок 2.1 – Діаграма варіантів використання системи")
    add_table(
        doc,
        "Таблиця 2.1 – Ролі користувачів і межі доступу в системі",
        ["Роль", "Доступні дії"],
        [
            ["Гість", "перегляд головної сторінки, каталогу, сторінок товарів"],
            ["Користувач", "обране, кошик, checkout, історія замовлень, сторінка оплати"],
            ["Адміністратор", "перегляд метрик, CRUD-операції з основними ресурсами"],
        ],
    )
    add_table(
        doc,
        "Таблиця 2.2 – Основні сценарії взаємодії із замовленням",
        ["Сценарій", "Ключові кроки"],
        [
            ["Створення", "додавання товарів у кошик, заповнення форми, створення order та payment_transaction"],
            ["Перегляд", "відкриття списку замовлень і деталізації конкретного замовлення"],
            ["Повторення", "перенесення позицій з попереднього замовлення до кошика"],
            ["Супровід оплати", "відображення статусу та синхронізація webhook-події"],
        ],
    )
    add_heading_block(doc, "2.2 Проєктування інформаційної моделі", 2)
    paragraph_block(doc, [
        "Інформаційна модель побудована на основі декількох ключових сутностей. Таблиця users використовується для зберігання облікових записів користувачів і зв’язується з кошиком, обраним і замовленнями. Таблиці categories і products описують каталог ювелірних виробів.",
        "Таблиця cart_items зберігає тимчасовий склад поточного кошика користувача. Після завершення checkout інформація з неї переноситься в orders та order_items, а записи кошика видаляються. Такий підхід забезпечує розмежування між чернетковим наміром купівлі та зафіксованим комерційним документом.",
        "Таблиця payment_transactions зберігає технічні атрибути платежу: провайдера, спосіб оплати, внутрішній і зовнішній референси, стан оплати, суму та службовий payload. Вона пов’язана один-до-одного із замовленням і дозволяє синхронізувати видимий стан замовлення з фактичним станом транзакції.",
    ])
    add_figure(doc, figures["2.2"], "Рисунок 2.2 – Діаграма бази даних системи", width_cm=15.2)
    add_table(
        doc,
        "Таблиця 2.3 – Основні сутності бази даних системи",
        ["Сутність", "Призначення"],
        [
            ["users", "зберігання облікових записів клієнтів"],
            ["categories", "довідник категорій ювелірних виробів"],
            ["products", "товарні позиції каталогу"],
            ["favorites", "обрані товари конкретного користувача"],
            ["cart_items", "поточні позиції в кошику"],
            ["orders", "заголовок замовлення та його агрегований стан"],
            ["order_items", "позиції конкретного замовлення"],
            ["payment_transactions", "технічний життєвий цикл оплати"],
        ],
    )
    add_heading_block(doc, "2.3 Проєктування поведінки системи", 2)
    paragraph_block(doc, [
        "Поведінка системи в основному обертається навколо двох процесів: оформлення замовлення та синхронізації оплати. У першому випадку система повинна перевірити склад кошика, сформувати агреговане замовлення, зберегти його позиції та створити пов’язану платіжну транзакцію. У другому випадку вона повинна прийняти повідомлення від платіжного провайдера, перевірити підпис webhook-повідомлення та узгодити статус замовлення з фактичним статусом оплати.",
        "Для опису цих сценаріїв використано діаграму діяльності процесу checkout і діаграму взаємодії при обробці webhook-повідомлення. Обидві діаграми показують, що бізнес-логіка реалізується не в контролерах, а на рівні сервісів, що відповідає принципу розділення відповідальності.",
    ])
    add_figure(doc, figures["2.3"], "Рисунок 2.3 – Діаграма діяльності процесу оформлення замовлення")
    add_figure(doc, figures["2.4"], "Рисунок 2.4 – Сценарій синхронізації оплати через webhook")

    add_heading_block(doc, "3 ПРОГРАМНА РЕАЛІЗАЦІЯ", 1)
    add_heading_block(doc, "3.1 Архітектурна концепція та модульна структура", 2)
    paragraph_block(doc, [
        "Архітектурно проєкт Diva Jewelry належить до класу сучасних монолітних вебзастосунків із вираженим розділенням на логічні модулі. Серверна частина реалізує маршрути, контролери, сервісний шар, моделі та політики доступу. Клієнтська частина побудована на Inertia-сторінках Vue, що дозволяє зберегти цілісну маршрутизацію без окремого API-окруження.",
        "Модульна структура проєкту відображає предметну декомпозицію. Окремо виділяються каталог, checkout, замовлення, платежі, обране, адміністративний контур і блок операційного моніторингу. У такий спосіб програмна структура безпосередньо відповідає логіці бізнес-процесів.",
    ])
    add_figure(doc, figures["3.1"], "Рисунок 3.1 – Діаграма розгортання системи Diva Jewelry")
    add_figure(doc, figures["3.2"], "Рисунок 3.2 – Діаграма компонентів програмної системи")
    add_table(
        doc,
        "Таблиця 3.1 – Архітектурні рівні системи",
        ["Рівень", "Відповідальність"],
        [
            ["Інтерфейс", "Vue-сторінки, форми, фільтри, відображення статусів"],
            ["HTTP-рівень", "маршрути, контролери, перевірка запитів, авторизація"],
            ["Сервісний шар", "checkout, каталог, webhook-обробка, синхронізація платежів"],
            ["Рівень даних", "Eloquent-моделі, міграції, зв’язки таблиць"],
            ["Інфраструктура", "Docker, Nginx, MySQL, Redis, queue, metrics"],
        ],
    )
    add_heading_block(doc, "3.2 Структура програмного забезпечення", 2)
    paragraph_block(doc, [
        "Поточний репозиторій містить 8 основних моделей, 24 контролери, 6 сервісів доменної логіки, 21 Vue-сторінку й 15 міграцій бази даних. Такий обсяг достатній для демонстрації повноцінної вебсистеми з користувацьким і службовим контурами.",
        "Каталог реалізовано через CatalogService, який централізує застосування пошуку, категорійних фільтрів, цінового діапазону, ознаки новинки та сортування. Сервіс також кешує список категорій і загальний діапазон цін, що зменшує повторні звернення до бази даних.",
        "Користувацькі сторінки каталогу, товару, checkout, історії замовлень і статусу оплати реалізовані у `resources/js/Pages`. Усі вони отримують дані безпосередньо через Inertia-відповіді контролерів, що дає змогу уникнути дублювання DTO-шарів для окремого клієнтського API.",
    ])
    add_table(
        doc,
        "Таблиця 3.2 – Технологічний стек системи",
        ["Технологія", "Роль у проєкті"],
        [
            ["Laravel 10", "серверна логіка, маршрути, ORM, авторизація"],
            ["Vue 3", "реактивний інтерфейс клієнтської частини"],
            ["Inertia.js", "зв’язок між Laravel і Vue-сторінками"],
            ["MySQL 8", "основне реляційне сховище"],
            ["Redis", "кешування й службові метрики"],
            ["Docker Compose", "контейнерне середовище запуску"],
        ],
    )
    add_figure(doc, figures["3.4"], "Рисунок 3.4 – Інтерфейс каталогу товарів")
    add_figure(doc, figures["3.5"], "Рисунок 3.5 – Сторінка перегляду окремого товару")
    add_heading_block(doc, "3.3 Реалізація створення та обробки замовлень", 2)
    paragraph_block(doc, [
        "Ключовим елементом реалізації є CheckoutService. Саме він відповідає за транзакційне створення замовлення на основі поточного кошика користувача. Під час виконання методу `createOrderFromCart` сервіс блокує записи кошика, перевіряє його коректність, визначає платіжний шлюз, створює запис order і пов’язану платіжну транзакцію, а потім формує позиції замовлення.",
        "Після успішного створення замовлення система очищає кошик користувача та скидає лічильник CartCounter. Додатково реєструється аудит-лог і збираються метрики про успішність або неуспішність checkout-процесу. Така побудова сервісу важлива, оскільки поєднує предметну логіку, технічну трасованість і контроль транзакційної цілісності.",
        "Сторінка оформлення замовлення в клієнтському інтерфейсі містить список позицій, форму контактних даних і вибір способу оплати. На відміну від небезпечних сценаріїв зберігання карткових реквізитів, реалізований варіант використовує лише демо-оплату та післяплату, не зберігаючи чутливих даних платіжної картки в базі.",
    ])
    add_figure(doc, figures["3.6"], "Рисунок 3.6 – Сторінка оформлення замовлення")
    add_table(
        doc,
        "Таблиця 3.3 – Відповідність сервісів і їх відповідальності",
        ["Сервіс", "Основна відповідальність"],
        [
            ["CatalogService", "пошук, фільтрація, сортування та кешування каталогу"],
            ["CheckoutService", "створення замовлення з кошика"],
            ["PaymentManager", "вибір платіжного шлюзу за методом оплати"],
            ["PaymentWebhookService", "перевірка та обробка webhook-повідомлення"],
            ["PaymentOrderSynchronizer", "узгодження статусу транзакції із замовленням"],
            ["PaymentReconciliationService", "службова звірка збережених транзакцій"],
        ],
    )
    add_table(
        doc,
        "Таблиця 3.4 – Етапи алгоритму створення замовлення",
        ["Етап", "Опис"],
        [
            ["1", "отримання і блокування позицій кошика користувача"],
            ["2", "валідація наявності товарів у кошику"],
            ["3", "створення агрегованого замовлення Order"],
            ["4", "ініціалізація PaymentTransaction через платіжний шлюз"],
            ["5", "збереження OrderItem для кожної позиції"],
            ["6", "очищення кошика та фіксація аудит-логу"],
        ],
    )
    add_heading_block(doc, "3.4 Реалізація платежів, станів і моніторингу", 2)
    paragraph_block(doc, [
        "Платіжний контур побудовано навколо інтерфейсу PaymentGateway і сервісу PaymentManager. Менеджер знаходить конкретну реалізацію шлюзу за назвою способу оплати. Це дозволяє легко додавати нові способи оплати без переписування checkout-логіки.",
        "Сервіс PaymentWebhookService відповідає за перевірку підпису webhook-повідомлення, розбір payload, пошук відповідної транзакції та її оновлення. Після цього PaymentOrderSynchronizer узгоджує статус транзакції зі станом замовлення. Якщо транзакція підтверджена як `paid`, замовлення також переходить у стан `paid`; якщо ж оплата неуспішна, стан замовлення синхронізується як `failed` або `cancelled`.",
        "Для технічного контролю застосовано власне сховище метрик MetricStore. Воно накопичує лічильники, gauge-значення та гістограми через Redis-backed кеш. Зокрема, у checkout-ланцюгу реєструються метрики тривалості обробки та кількості спроб, а в webhook-обробнику – метрики вхідних запитів за провайдером та результатом валідації.",
    ])
    add_table(
        doc,
        "Таблиця 3.5 – Аналітичні показники та службові сигнали системи",
        ["Показник", "Призначення"],
        [
            ["checkout_orders_total", "кількість спроб checkout за результатом і методом оплати"],
            ["checkout_duration_seconds", "тривалість обробки checkout"],
            ["payment_webhooks_total", "кількість webhook-подій за результатом перевірки"],
            ["latestOrders", "зріз останніх транзакцій для бек-офісу"],
        ],
    )
    add_figure(doc, figures["3.7"], "Рисунок 3.7 – Сторінка статусу оплати")
    add_heading_block(doc, "3.5 Адміністративна панель роботи із замовленнями", 2)
    paragraph_block(doc, [
        "Адміністративний контур реалізовано у вигляді окремого маршрутизованого блоку з формою входу, дашбордом і універсальним ресурсним контролером. На дашборді виводяться метрики по товарах, категоріях, клієнтах, замовленнях, платежах і обраних позиціях, а також список останніх замовлень із сумою та статусами.",
        "Контролер ResourceController реалізує універсальний CRUD-підхід. Він отримує конфігурацію ресурсу з реєстру, формує схему колонок і фільтрів, перевіряє дозволи й віддає уніфіковані сторінки списку та форми. Такий підхід зменшує дублювання коду між товарами, категоріями, користувачами, замовленнями та платежами.",
        "Особливу цінність для задачі обліку має можливість бачити разом як бізнесові, так і технічні поля замовлення. Адміністратор отримує доступ до загального статусу замовлення, статусу оплати, платіжного референсу, суми та пов’язаних записів. Це суттєво скорочує час пошуку та узгодження транзакцій.",
    ])
    add_figure(doc, figures["3.8"], "Рисунок 3.8 – Панель адміністратора: аналітичний блок")
    add_table(
        doc,
        "Таблиця 3.6 – Адміністративні модулі, пов’язані з життєвим циклом замовлення",
        ["Модуль", "Роль у процесі"],
        [
            ["Товари", "підтримка номенклатури й цін"],
            ["Категорії", "групування виробів для каталогу та головної сторінки"],
            ["Замовлення", "облік агрегованих покупок і контроль станів"],
            ["Платежі", "аналіз технічних параметрів транзакції"],
            ["Користувачі", "пошук власника замовлення та контексту покупки"],
            ["Обране", "маркетинговий сигнал за інтересом до товарів"],
        ],
    )
    add_heading_block(doc, "3.6 Контейнеризація, запуск і супровід середовища", 2)
    paragraph_block(doc, [
        "Середовище розгортання описано у файлі `docker-compose.yml`. Воно містить сервіси `web`, `app`, `vite`, `queue`, `db` і `redis`. Таким чином, окремо виділяються вебсервер, PHP-застосунок, HMR-оточення для фронтенду, асинхронний worker, база даних і Redis.",
        "Контейнеризація вирішує одразу декілька задач. По-перше, забезпечує однакове середовище для розробки та демонстрації. По-друге, спрощує залежності: локальній машині не потрібні окремо встановлені PHP, MySQL чи Redis. По-третє, дозволяє явно описати мережеву взаємодію між компонентами та політику їхнього запуску.",
        "Для контролю працездатності застосовано окремі маршрути `/live`, `/ready`, `/up` і `/metrics`. Це важливо для дипломного проєкту, оскільки демонструє не лише функціональний контур продажу, а й базову операційну зрілість системи.",
    ])
    add_table(
        doc,
        "Таблиця 3.7 – Інфраструктурні сервіси середовища розгортання",
        ["Сервіс", "Призначення"],
        [
            ["web", "проксіювання HTTP-запитів через Nginx"],
            ["app", "виконання Laravel-застосунку на PHP-FPM"],
            ["vite", "фронтенд-розробка та HMR"],
            ["queue", "фоновий worker для службових задач"],
            ["db", "сховище даних MySQL 8"],
            ["redis", "кеш і службова транспортна шина"],
        ],
    )

    add_heading_block(doc, "4 ТЕСТУВАННЯ", 1)
    add_heading_block(doc, "4.1 Підхід до тестування та критерії оцінювання", 2)
    paragraph_block(doc, [
        "Тестування системи виконувалося на двох рівнях: функціональному та автоматизованому. Функціональний рівень перевіряє завершені користувацькі сценарії, а автоматизований – окремі технічні інваріанти бізнес-логіки на стороні Laravel.",
        "Критеріями оцінювання були коректність побудови маршруту користувача, узгодженість станів замовлення й оплати, правильне очищення кошика після checkout, доступність історії замовлень та працездатність адміністративних ресурсів.",
    ])
    add_heading_block(doc, "4.2 Сценарії функціонального тестування", 2)
    paragraph_block(doc, [
        "Функціональне тестування охоплює перевірку каталогу, сторінки товару, кошика, оформлення замовлення, сторінки статусу оплати, історії замовлень і бек-офісу. Для кожного сценарію важливо оцінити як позитивний шлях, так і реакцію системи на невалідний стан, наприклад порожній кошик або відсутність дозволів.",
    ])
    add_figure(doc, figures["4.1"], "Рисунок 4.1 – Приклад сторінки замовлень")
    add_table(
        doc,
        "Таблиця 4.1 – Основні тестові сценарії системи обліку замовлень",
        ["Сценарій", "Очікуваний результат"],
        [
            ["Пошук у каталозі", "список товарів фільтрується без повного перезавантаження"],
            ["Додавання в кошик", "товар з’являється в кошику з коректною кількістю"],
            ["Checkout", "створюються order, order_items і payment_transaction"],
            ["Webhook paid", "замовлення та транзакція переходять у стан paid"],
            ["Повторне замовлення", "позиції повертаються до кошика поточного користувача"],
            ["CRUD у бек-офісі", "запис створюється, редагується й видаляється за наявності прав"],
        ],
    )
    add_heading_block(doc, "4.3 Аналіз отриманих результатів", 2)
    paragraph_block(doc, [
        "У результаті перевірки встановлено, що система стабільно проходить основні користувацькі маршрути. Каталог коректно застосовує фільтри, checkout формує цілісний агрегат замовлення, а механізм webhook-симуляції дозволяє наочно перевірити синхронізацію платіжного стану без реального еквайрингу.",
        "Також підтверджено коректність бек-офісного контуру. Дашборд відображає агреговані метрики, а універсальний ресурсний контролер забезпечує предиктивну роботу з таблицями даних. Це свідчить про достатню узгодженість між користувацькою і службовою частинами системи.",
    ])
    add_heading_block(doc, "4.4 Автоматизоване тестування Laravel", 2)
    paragraph_block(doc, [
        "Автоматизоване тестування в проєкті ґрунтується на Laravel Test Suite. У репозиторії присутні тести, орієнтовані на перевірку сервісів кошика й checkout-логіки. Для тестового оточення використовується SQLite in-memory, що прискорює виконання перевірок і ізолює тести від продуктивної конфігурації MySQL.",
        "Для задачі дипломної роботи такий підхід є достатнім, оскільки він підтверджує працездатність найризикованішої бізнес-логіки: створення замовлення, очищення кошика та супроводу платіжного контуру. Надалі набір тестів може бути розширений перевірками адміністрування ресурсів та метрик.",
    ])

    add_heading_block(doc, "5 ОРГАНІЗАЦІЙНЕ ЗАБЕЗПЕЧЕННЯ", 1)
    add_heading_block(doc, "5.1 Організація експлуатації системи", 2)
    paragraph_block(doc, [
        "Організаційне забезпечення експлуатації вебсистеми включає правила взаємодії між покупцем, менеджером, адміністратором і технічним персоналом. На практиці це означає, що система повинна вписуватися в щоденний цикл роботи магазину: оновлення номенклатури, обробку нових замовлень, контроль оплат, підтримку довідників і технічний моніторинг.",
        "Для стабільної експлуатації важливо закріпити межі відповідальності між ролями. Покупець працює лише з власними даними та покупками, менеджер контролює комерційні операції, адміністратор відповідає за модель доступу та довідники, а технічний спеціаліст підтримує середовище запуску, резервні копії та контроль health-check маршрутів.",
    ])
    add_table(
        doc,
        "Таблиця 5.1 – Основні організаційні дії під час експлуатації системи",
        ["Дія", "Періодичність"],
        [
            ["перевірка нових замовлень", "протягом робочого дня"],
            ["контроль статусів оплати", "після кожної платіжної події"],
            ["актуалізація каталогу", "за потребою або за графіком контент-оновлень"],
            ["перевірка health-check і metrics", "щоденно"],
            ["резервне копіювання даних", "регламентовано політикою супроводу"],
        ],
    )
    add_heading_block(doc, "5.1.1 Функціонування системи для ролі користувач", 2)
    paragraph_block(doc, [
        "Для ролі користувача система забезпечує послідовний маршрут: вибір категорії, робота з каталогом, додавання виробів в обране, керування кошиком і checkout. Після оформлення замовлення користувач переходить на сторінку статусу оплати, а в подальшому може повернутися до історії замовлень та повторити покупку.",
        "Такий сценарій є достатньо компактним і зрозумілим, що зменшує кількість помилок при взаємодії з інтерфейсом. Водночас він підтримує всі ключові операції предметної області продажу ювелірних виробів.",
    ])
    add_heading_block(doc, "5.1.2 Функціонування системи для ролі адміністратор", 2)
    paragraph_block(doc, [
        "Для ролі адміністратора система концентрує дані в межах єдиного бек-офісу. На стартовій сторінці доступні швидкі метрики по товарах, клієнтах, замовленнях і платежах. Далі адміністратор переходить до ресурсних таблиць, де виконує створення, редагування та видалення записів.",
        "Ключовою властивістю адміністративного контуру є уніфікація інтерфейсу. Незалежно від ресурсу, адміністратор працює з одними й тими самими базовими діями: пошук, сортування, фільтрація, створення та редагування. Це знижує поріг входу для супроводу системи.",
    ])
    add_heading_block(doc, "5.2 Розподіл ролей, доступу та супроводу", 2)
    paragraph_block(doc, [
        "Розподіл ролей повинен бути формалізований як на рівні організації робіт, так і на рівні програмного доступу. У проєкті це підтримується через окремі middleware, політики й механізми бек-офісної авторизації. Таким чином, користувач не може отримати доступ до службового контуру, а адміністратор працює в окремому захищеному середовищі.",
        "З погляду супроводу важливо контролювати не лише бізнесові дані, а й технічні журнали, аудит-логи та метрики. Для реальної експлуатації це означає регулярний перегляд стану контейнерів, бази даних, Redis-сервісу та накопичених транзакцій, що не були остаточно узгоджені.",
    ])
    add_figure(doc, figures["5.1"], "Рисунок 5.1 – Розподіл ролей під час експлуатації")
    add_figure(doc, figures["5.2"], "Рисунок 5.2 – Сторінка універсального CRUD-ресурсу")
    add_table(
        doc,
        "Таблиця 5.2 – Розподіл відповідальності між учасниками експлуатації",
        ["Роль", "Відповідальність"],
        [
            ["Покупець", "оформлення й контроль власних замовлень"],
            ["Менеджер", "операційний супровід замовлень і контенту"],
            ["Адміністратор", "керування ресурсами та доступами"],
            ["DevOps/технічний супровід", "контейнери, моніторинг, резервування"],
        ],
    )

    add_heading_block(doc, "ВИСНОВКИ", 1)
    paragraph_block(doc, [
        "У дипломній роботі було розглянуто, спроєктовано та описано веборієнтовану інформаційну систему обліку замовлення та продажу ювелірних виробів Diva Jewelry.",
        "У першому розділі проаналізовано предметну область продажу ювелірних виробів і обґрунтовано необхідність централізованого обліку замовлень як ключової бізнес-сутності електронної комерції.",
        "У другому розділі виконано проєктування користувацьких сценаріїв, реляційної моделі даних та основних поведінкових процесів системи, зокрема checkout і синхронізації оплати через webhook.",
        "У третьому розділі описано програмну реалізацію системи на основі Laravel, Vue 3, Inertia.js, MySQL, Redis та Docker. Показано, що сервіси CheckoutService, PaymentWebhookService і PaymentOrderSynchronizer утворюють послідовний механізм створення й супроводу замовлення.",
        "У четвертому розділі розглянуто тестування системи. Перевірено основні користувацькі та адміністративні сценарії, а також підтверджено працездатність автоматизованого тестового контуру.",
        "У п’ятому розділі визначено організаційні умови експлуатації системи, розподіл ролей і базові вимоги до технічного супроводу, моніторингу та резервування.",
        "Отже, поставлену мету дипломної роботи досягнуто. Розроблена система може використовуватися як основа для подальшого розвитку комерційної платформи продажу ювелірних виробів, зокрема для інтеграції реальних платіжних шлюзів, розширення аналітики та поглиблення механізмів керування товарними залишками.",
    ])

    add_heading_block(doc, "ПЕРЕЛІК ДЖЕРЕЛ", 1)
    for item in [
        "1. ДСТУ 8302:2015. Інформація та документація. Бібліографічне посилання. Загальні положення та правила складання. Київ : ДП «УкрНДНЦ», 2016. 16 с.",
        "2. ISO/IEC 25010:2011. Systems and software engineering — Systems and software quality models. Geneva : ISO, 2011.",
        "3. Sommerville I. Software Engineering. 10th ed. Boston : Pearson, 2015. 816 p.",
        "4. Pressman R. S., Maxim B. R. Software Engineering: A Practitioner’s Approach. 9th ed. New York : McGraw-Hill, 2020. 704 p.",
        "5. Fowler M. Patterns of Enterprise Application Architecture. Boston : Addison-Wesley, 2002. 560 p.",
        "6. Fowler M. UML Distilled. 3rd ed. Boston : Addison-Wesley, 2003. 208 p.",
        "7. Larman C. Applying UML and Patterns. 3rd ed. Upper Saddle River : Prentice Hall, 2004. 736 p.",
        "8. Elmasri R., Navathe S. Fundamentals of Database Systems. 7th ed. Boston : Pearson, 2016. 1280 p.",
        "9. Connolly T., Begg C. Database Systems. 6th ed. Boston : Pearson, 2015. 1440 p.",
        "10. Gamma E., Helm R., Johnson R., Vlissides J. Design Patterns. Boston : Addison-Wesley, 1994. 395 p.",
        "11. Laravel. Laravel Documentation. URL: https://laravel.com/docs (дата звернення: 17.05.2026).",
        "12. Laravel. Eloquent ORM Documentation. URL: https://laravel.com/docs/eloquent (дата звернення: 17.05.2026).",
        "13. Laravel. Testing Documentation. URL: https://laravel.com/docs/testing (дата звернення: 17.05.2026).",
        "14. PHP Group. PHP Manual. URL: https://www.php.net/manual/en/ (дата звернення: 17.05.2026).",
        "15. Vue.js. Vue.js Guide. URL: https://vuejs.org/guide/introduction.html (дата звернення: 17.05.2026).",
        "16. Inertia.js. Inertia.js Documentation. URL: https://inertiajs.com/ (дата звернення: 17.05.2026).",
        "17. Vite. Vite Guide. URL: https://vite.dev/guide/ (дата звернення: 17.05.2026).",
        "18. MySQL. MySQL 8.0 Reference Manual. URL: https://dev.mysql.com/doc/refman/8.0/en/ (дата звернення: 17.05.2026).",
        "19. Redis. Redis Documentation. URL: https://redis.io/docs/latest/ (дата звернення: 17.05.2026).",
        "20. Docker. Docker Documentation. URL: https://docs.docker.com/ (дата звернення: 17.05.2026).",
        "21. Docker. Docker Compose Documentation. URL: https://docs.docker.com/compose/ (дата звернення: 17.05.2026).",
        "22. NGINX. Documentation. URL: https://nginx.org/en/docs/ (дата звернення: 17.05.2026).",
        "23. PHPUnit. Documentation. URL: https://phpunit.de/documentation.html (дата звернення: 17.05.2026).",
        "24. OWASP Foundation. OWASP Top 10. URL: https://owasp.org/www-project-top-ten/ (дата звернення: 17.05.2026).",
    ]:
        add_paragraph(doc, item)

    add_heading_block(doc, "ДОДАТКИ", 1)
    add_heading_block(doc, "ДОДАТОК А – Декларація про використання систем ШІ", 1)
    paragraph_block(doc, [
        "Інструментарій: OpenAI ChatGPT, модель GPT-5.5.",
        "Локалізація та обсяг: система ШІ використовувалася для підготовки чернетки пояснювальної записки, стилістичного редагування окремих фрагментів, формування текстових пояснень до таблиць і рисунків, а також для структуризації аналітичного викладу. Архітектурні висновки, змістове наповнення, зв’язок із кодовою базою проєкту й остаточна перевірка документа виконувалися автором.",
        "Методологія використання: система ШІ застосовувалася як допоміжний інструмент для прискорення підготовки тексту, уніфікації термінології та перевірки зв’язності викладу. ШІ не використовувався як заміна програмної реалізації або фактичних результатів аналізу репозиторію.",
        "Період роботи: травень 2026 року.",
    ])

    add_heading_block(doc, "ДОДАТОК Б – Структура проєкту", 1)
    add_paragraph(
        doc,
        "├── app\n│   ├── Http\n│   │   ├── Controllers\n│   │   ├── Middleware\n│   │   └── Requests\n│   ├── Models\n│   │   ├── Category.php\n│   │   ├── Product.php\n│   │   ├── Order.php\n│   │   ├── OrderItem.php\n│   │   ├── PaymentTransaction.php\n│   │   └── User.php\n│   ├── Services\n│   │   ├── CatalogService.php\n│   │   ├── CheckoutService.php\n│   │   ├── PaymentManager.php\n│   │   ├── PaymentOrderSynchronizer.php\n│   │   ├── PaymentReconciliationService.php\n│   │   └── PaymentWebhookService.php\n├── database\n│   ├── factories\n│   ├── migrations\n│   └── seeders\n├── resources\n│   ├── js\n│   │   ├── Components\n│   │   ├── Layouts\n│   │   └── Pages\n│   └── views\n├── routes\n│   ├── web.php\n│   └── api.php\n├── docker-compose.yml\n├── Dockerfile\n├── composer.json\n└── package.json",
    )

    add_heading_block(doc, "ДОДАТОК В – Фрагмент сервісу створення замовлення", 1)
    add_paragraph(
        doc,
        "public function createOrderFromCart(User $user, array $payload): Order\n{\n    return DB::transaction(function () use ($user, $payload): Order {\n        $items = CartItem::query()\n            ->where('user_id', $user->id)\n            ->with('product')\n            ->lockForUpdate()\n            ->get();\n\n        $this->guardAgainstInvalidCart($items);\n        $gateway = $this->paymentManager->gatewayForMethod((string) $payload['payment_method']);\n\n        $order = Order::create([\n            'user_id' => $user->id,\n            'full_name' => $payload['full_name'],\n            'email' => $payload['email'],\n            'payment_method' => $payload['payment_method'],\n            'payment_provider' => $gateway->key(),\n            'payment_reference' => $this->generatePaymentReference(),\n            'payment_status' => 'pending',\n            'total' => $items->sum(fn (CartItem $item) => $item->product->price * $item->quantity),\n            'status' => 'pending',\n        ]);\n    });\n}",
    )

    add_heading_block(doc, "ДОДАТОК Г – Інтерфейси користувача", 1)
    add_figure(doc, figures["G.1"], "Рисунок Г.1 – Сторінка деталізації замовлення користувача")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
